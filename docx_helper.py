import zipfile, os, re, io
from docx import Document, shared


class DocxHelper:
    def __init__(self, input_file, output_file, yearly_data, profile_data):
        self.input_file = input_file
        self.output_file = output_file
        self.yearly_data = yearly_data
        self.profile_data = profile_data
    
    def replace_text(self, text):
        matches = re.findall(r"<([^<>]+)>", text)
        for match in matches:
            
            if ("Image" not in match):
                if ("|" in match):
                    filter = match.split("|")[1]
                    tagname = match.split("|")[0]
                    if (tagname in self.yearly_data[filter]):
                        text = text.replace("<" + match + ">", str(self.yearly_data[filter][tagname]))
                else:
                    if (match in self.profile_data):
                        text = text.replace("<" + match + ">", str(self.profile_data[match]))
                    
        return text

    def replace_images(self, text, r):
        image_tags = re.findall(r"<Image:[^>]+>", text)
        
        for tag in image_tags:
            parts = tag[1:-1].split(':')
            
            if len(parts) >= 3:
                tag_name = parts[1]
                
                try:
                    width = float(parts[2])
                    width_emu = width * 12700
                    
                    image_data = self.profile_data["Images"].get(tag_name)
                    
                    if isinstance(image_data, bytes):
                        image_stream = io.BytesIO(image_data)
                        r.add_picture(image_stream, width=width_emu)
                        
                    text = text.replace(tag, "")
                except ValueError:
                    pass
        
        return text

    def generate_docx(self):
        document = Document(self.input_file)

        # Replace text in tables
        for table in document.tables:
            for column in table.columns:
                for cell in column.cells:
                    if cell.paragraphs:
                        for p in cell.paragraphs:
                            full_text = "".join(run.text for run in p.runs)  # Get full paragraph text
                            full_text = self.replace_text(full_text)
                            font_props = None
                            if p.runs:
                                run = p.runs[0]
                                font_props = {
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'size': run.font.size,
                                    'name': run.font.name,
                                    'color': run.font.color.rgb if run.font.color else None
                                }
                            p.clear()
                            new_run = p.add_run(full_text)
                            
                            # Apply stored formatting
                            if font_props:
                                new_run.bold = font_props['bold']
                                new_run.italic = font_props['italic']
                                new_run.font.size = font_props['size']
                                new_run.font.name = font_props['name']
                                if font_props['color']:
                                    new_run.font.color.rgb = font_props['color']
                            
        
        # Replace text in paragraphs

        for p in document.paragraphs:
            full_text = "".join(run.text for run in p.runs)  # Get full paragraph text
            full_text = self.replace_text(full_text)

            if p.runs:
                font_size = p.runs[0].font.size
            p.clear()

            image_run = p.add_run()
            full_text = self.replace_images(full_text, image_run)
            run = p.add_run(full_text)

            if (font_size):
                run.font.size = font_size

        document.save(self.output_file)
        return os.path.abspath(self.output_file)